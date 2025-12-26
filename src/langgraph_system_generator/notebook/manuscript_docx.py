"""Professional DOCX manuscript generation for print outputs."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


class ManuscriptDOCXGenerator:
    """Generates formatted DOCX manuscripts with professional styling."""

    def __init__(
        self,
        font_name: str = "Times New Roman",
        font_size: int = 12,
        line_spacing: float = 2.0,
    ):
        """Initialize the DOCX generator with default styling.

        Args:
            font_name: Font family for body text.
            font_size: Font size in points for body text.
            line_spacing: Line spacing multiplier (e.g., 2.0 for double-spacing).
        """
        self.font_name = font_name
        self.font_size = font_size
        self.line_spacing = line_spacing

    def create_manuscript(
        self,
        title: str,
        author: str | None = None,
        chapters: Sequence[Dict[str, Any]] | None = None,
        output_path: str | Path = "manuscript.docx",
        include_title_page: bool = True,
    ) -> str:
        """Generate a print-ready DOCX manuscript.

        Args:
            title: Manuscript title.
            author: Author name (optional).
            chapters: List of chapter dictionaries with 'title' and 'content' keys.
                     Content can be a string or list of strings.
            output_path: Destination path for the DOCX file.
            include_title_page: Whether to include a formatted title page.

        Returns:
            Path to the created DOCX file.
        """
        target = Path(output_path)
        target.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._configure_styles(doc)

        if include_title_page and title:
            self._add_title_page(doc, title, author)

        if chapters:
            for chapter in chapters:
                self._add_chapter(doc, chapter)

        doc.save(str(target))
        return str(target)

    def _configure_styles(self, doc: Document) -> None:
        """Configure document-wide styles for professional appearance.

        Args:
            doc: The Document object to style.
        """
        # Configure chapter title style (Heading 1)
        if "Heading 1" in doc.styles:
            chapter_style = doc.styles["Heading 1"]
            chapter_style.font.name = self.font_name
            chapter_style.font.size = Pt(16)
            chapter_style.font.bold = True
            chapter_style.font.color.rgb = RGBColor(0, 0, 0)

        # Configure section heading style (Heading 2)
        if "Heading 2" in doc.styles:
            section_style = doc.styles["Heading 2"]
            section_style.font.name = self.font_name
            section_style.font.size = Pt(14)
            section_style.font.bold = True

        # Configure body text style
        if "Normal" in doc.styles:
            normal_style = doc.styles["Normal"]
            normal_style.font.name = self.font_name
            normal_style.font.size = Pt(self.font_size)
            normal_style.paragraph_format.line_spacing = self.line_spacing
            normal_style.paragraph_format.space_after = Pt(6)

    def _add_title_page(self, doc: Document, title: str, author: str | None = None) -> None:
        """Add a formatted title page to the document.

        Args:
            doc: The Document object.
            title: The manuscript title.
            author: The author name (optional).
        """
        # Add title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.name = self.font_name
        title_run.font.size = Pt(24)
        title_run.font.bold = True

        # Add spacing
        for _ in range(3):
            doc.add_paragraph()

        # Add author if provided
        if author:
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"by {author}")
            author_run.font.name = self.font_name
            author_run.font.size = Pt(14)

        # Add page break after title page
        doc.add_page_break()

    def _add_chapter(self, doc: Document, chapter: Dict[str, Any]) -> None:
        """Add a chapter to the document.

        Args:
            doc: The Document object.
            chapter: Dictionary with 'title' and 'content' keys.
                    Content can be a string, list of strings, or list of paragraphs.
        """
        chapter_title = chapter.get("title", "Untitled Chapter")
        chapter_content = chapter.get("content", [])

        # Add chapter title
        doc.add_heading(chapter_title, level=1)

        # Add chapter content
        if isinstance(chapter_content, str):
            # Single string content
            for paragraph in chapter_content.split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
        elif isinstance(chapter_content, (list, tuple)):
            # List of paragraphs or sections
            for item in chapter_content:
                if isinstance(item, dict):
                    # Structured content with sections
                    section_title = item.get("heading")
                    section_text = item.get("text", "")

                    if section_title:
                        doc.add_heading(section_title, level=2)

                    if section_text:
                        doc.add_paragraph(section_text)
                else:
                    # Plain paragraph text
                    if item and str(item).strip():
                        doc.add_paragraph(str(item).strip())

        # Add page break after chapter
        doc.add_page_break()

    def create_notebook_manuscript(
        self,
        notebook_cells: Sequence[Dict[str, Any]],
        output_path: str | Path,
        title: str | None = None,
        author: str | None = None,
    ) -> str:
        """Create a manuscript from notebook cells.

        Args:
            notebook_cells: List of cell dictionaries with 'cell_type', 'content', and optional 'section'.
            output_path: Destination path for the DOCX file.
            title: Optional manuscript title.
            author: Optional author name.

        Returns:
            Path to the created DOCX file.
        """
        target = Path(output_path)
        target.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._configure_styles(doc)

        if title:
            self._add_title_page(doc, title, author)

        current_section = None

        for cell in notebook_cells:
            cell_type = cell.get("cell_type", "code")
            content = cell.get("content", "")
            section = cell.get("section")

            # Add section heading if changed
            if section and section != current_section:
                doc.add_heading(section.replace("_", " ").title(), level=1)
                current_section = section

            # Process cell content
            if cell_type == "markdown" and content.strip():
                # Parse markdown headings
                for line in content.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith("### "):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith("## "):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith("# "):
                        doc.add_heading(line[2:], level=1)
                    else:
                        doc.add_paragraph(line)

            elif cell_type == "code" and content.strip():
                # Add code as preformatted block
                code_para = doc.add_paragraph(content, style="Intense Quote")
                code_para.paragraph_format.left_indent = Inches(0.5)

        doc.save(str(target))
        return str(target)
