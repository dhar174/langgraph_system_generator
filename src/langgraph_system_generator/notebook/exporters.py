"""Utilities to export notebooks into distributable formats."""

from __future__ import annotations

import io
import os
import subprocess
import zipfile
from pathlib import Path
from typing import Sequence

import nbformat

# Base output directory for all notebook exports. This mirrors the API server's
# restriction and ensures exporters cannot write outside the configured root.
_BASE_OUTPUT = Path(os.environ.get("LNF_OUTPUT_BASE", ".")).resolve()


def _safe_output_path(path: str | os.PathLike[str]) -> Path:
    """Resolve an output path and ensure it stays within the allowed base directory.

    This provides a defense-in-depth guard so that exporters cannot be used to
    write files outside of the configured root directory, even when called
    directly from external code.
    """
    target = Path(path).resolve()
    output_dir = target.parent
    if not output_dir.is_relative_to(_BASE_OUTPUT):
        raise RuntimeError("Output directory must reside within the allowed base directory.")
    output_dir.mkdir(parents=True, exist_ok=True)
    return target


class NotebookExporter:
    """Exports nbformat notebooks to files or bundles."""

    def export_ipynb(self, notebook: nbformat.NotebookNode, path: str | Path) -> str:
        """Write a validated notebook to disk."""
        nbformat.validate(notebook)
        target = _safe_output_path(path)
        with target.open("w", encoding="utf-8") as handle:
            nbformat.write(notebook, handle)
        return str(target)

    def export_zip(
        self,
        notebook: nbformat.NotebookNode,
        zip_path: str | Path,
        extra_files: Sequence[str | os.PathLike[str]] | None = None,
        notebook_name: str = "notebook.ipynb",
    ) -> str:
        """Create a zip bundle containing the notebook and optional artifacts."""
        nbformat.validate(notebook)
        buffer = io.StringIO()
        nbformat.write(notebook, buffer)

        target = _safe_output_path(zip_path)
        with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(notebook_name, buffer.getvalue())
            for extra in extra_files or []:
                extra_path = Path(extra)
                if extra_path.is_file():
                    zf.write(extra_path, arcname=extra_path.name)
        return str(target)

    def export_to_html(
        self, notebook: nbformat.NotebookNode, output_path: str | Path
    ) -> str:
        """Export notebook to HTML using nbconvert.

        Args:
            notebook: The notebook to export.
            output_path: Destination path for the HTML file.

        Returns:
            Path to the created HTML file.

        Raises:
            ImportError: If nbconvert is not available.
            Exception: If export fails.
        """
        try:
            from nbconvert import HTMLExporter
        except ImportError as exc:
            raise ImportError(
                "nbconvert is required for HTML export. Install it with: pip install nbconvert"
            ) from exc

        nbformat.validate(notebook)
        target = _safe_output_path(output_path)

        exporter = HTMLExporter()
        (body, resources) = exporter.from_notebook_node(notebook)

        with target.open("w", encoding="utf-8") as handle:
            handle.write(body)

        return str(target)

    def export_to_pdf(
        self, notebook_path: str | Path, output_path: str | Path, method: str = "webpdf"
    ) -> str:
        """Export notebook to PDF using nbconvert.

        Args:
            notebook_path: Path to the source notebook file.
            output_path: Destination path for the PDF file.
            method: PDF export method - 'webpdf' (default, more reliable) or 'latex' (requires LaTeX).

        Returns:
            Path to the created PDF file.

        Raises:
            ImportError: If nbconvert is not available.
            RuntimeError: If export fails.
        """
        try:
            from nbconvert import PDFExporter
        except ImportError as exc:
            raise ImportError(
                "nbconvert is required for PDF export. Install it with: pip install nbconvert"
            ) from exc

        source = Path(notebook_path)
        if not source.exists():
            raise FileNotFoundError(f"Notebook not found: {source}")

        # Resolve the output path and ensure it stays within the allowed base directory.
        target = _safe_output_path(output_path)

        if method == "latex":
            # Use LaTeX-based PDF export (requires LaTeX installation)
            try:
                exporter = PDFExporter()
                (body, resources) = exporter.from_filename(str(source))

                with target.open("wb") as handle:
                    handle.write(body)
                return str(target)
            except Exception as exc:
                raise RuntimeError(
                    f"LaTeX-based PDF export failed: {exc}. "
                    "Try 'webpdf' method or ensure LaTeX is installed."
                ) from exc
        else:
            # Use webpdf method (more reliable, uses Chromium)
            try:
                # Note: nbconvert appends .pdf extension if not present, and if present it might double it depending on version/config
                # We specify output filename without extension if we want it to just append .pdf, or with extension.
                # But nbconvert behavior with --output is tricky.
                # If we pass --output /path/to/notebook.pdf, it might write to /path/to/notebook.pdf.pdf

                # Let's try to let nbconvert determine the output filename by specifying output-dir and output base name
                output_base = target.stem

                # Check if target ends with .pdf
                if target.suffix == '.pdf':
                    # If we explicitly want .pdf, we should be careful.
                    # nbconvert automatically adds the extension for the format.
                    pass

                cmd = [
                        "jupyter",
                        "nbconvert",
                        "--to",
                        "webpdf",
                        "--output-dir",
                        str(output_dir),
                        "--output",
                        output_base,
                        str(source.resolve()),
                    ]

                _ = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    check=True,
                )

                # Expected output file
                expected_output = output_dir / f"{output_base}.pdf"

                if expected_output.exists():
                    if expected_output != target:
                        # Rename if necessary (though if target was .pdf, expected_output should match)
                        if target.exists():
                            target.unlink()
                        expected_output.rename(target)
                    return str(target)
                else:
                    # Fallback check if it did something else
                     raise RuntimeError(f"PDF export finished but file not found at {expected_output}")

            except subprocess.CalledProcessError as exc:
                raise RuntimeError(
                    f"webpdf export failed: {exc.stderr}. "
                    "Ensure Jupyter and Chromium/Chrome are installed."
                ) from exc
            except FileNotFoundError as exc:
                raise RuntimeError(
                    "jupyter command not found. Ensure Jupyter is installed and in PATH."
                ) from exc

    def export_notebook_to_docx(
        self,
        notebook: nbformat.NotebookNode,
        output_path: str | Path,
        title: str | None = None,
    ) -> str:
        """Export notebook content to a DOCX document.

        This provides basic DOCX export. For more sophisticated formatting,
        use the ManuscriptDOCXGenerator class.

        Args:
            notebook: The notebook to export.
            output_path: Destination path for the DOCX file.
            title: Optional title for the document.

        Returns:
            Path to the created DOCX file.

        Raises:
            ImportError: If python-docx is not available.
        """
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX export. Install it with: pip install python-docx"
            ) from exc

        nbformat.validate(notebook)
        target = Path(output_path)
        target.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        if title:
            doc.add_heading(title, 0)

        for cell in notebook.cells:
            if cell.cell_type == "markdown":
                # Add markdown content as paragraphs
                for line in cell.source.split("\n"):
                    if line.strip():
                        if line.startswith("# "):
                            doc.add_heading(line[2:], level=1)
                        elif line.startswith("## "):
                            doc.add_heading(line[3:], level=2)
                        elif line.startswith("### "):
                            doc.add_heading(line[4:], level=3)
                        else:
                            doc.add_paragraph(line)
            elif cell.cell_type == "code":
                # Add code as preformatted text
                if cell.source.strip():
                    doc.add_paragraph(cell.source, style="Intense Quote")

        doc.save(str(target))
        return str(target)
